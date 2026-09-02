#!/usr/bin/env python3
"""
Render an SOP as the Word document a customer would be handed. Figures are
recovered by matching each step back to the video it came from, so every caption
names the run and second it shows.

    python pipeline/06_sop_to_docx.py [--group molly] [--no-images]
"""

import argparse
import json
import re
import sys
import tempfile
from pathlib import Path

INVESTIGATION = Path(__file__).resolve().parent.parent   # investigation_operator/
REPO = INVESTIGATION.parent                              # repository root
SOP_DIR = INVESTIGATION / "sops"
OUTPUT_DIR = INVESTIGATION / "output"
FRAME_DIR = INVESTIGATION / "frames"
# The .docx is handed to a customer, not joined to anything.
DELIVERABLE_DIR = REPO / "deliverables"

sys.path.insert(0, str(INVESTIGATION))
from _text import keywords, overlap  # noqa: E402

GROUP_RUNS = {"molly": ["molly1", "molly2", "molly3"],
              "baseline": ["alice", "georgia", "tori"]}

# Frames the automatic pick got wrong, corrected after watching the footage.
# An integer shifts the frame that many seconds, None drops the figure.
FIGURE_OVERRIDES = {
    # "the apple's stem sits in the path of a straight cut"
    ("molly3", 6): 1,        # the adjustment is visible a beat later
    # "the cut slices are spreading unevenly across the board"
    ("molly1", 48): None,    # not a legible moment on camera
}

# Filled in as overrides fire, so main() can flag any that never matched.
_overrides_used: set[tuple[str, int]] = set()


def override_for(video_id: str, ts: int):
    """Return (shift, was_overridden) for one candidate frame."""
    key = (video_id, ts)
    if key in FIGURE_OVERRIDES:
        _overrides_used.add(key)
        return FIGURE_OVERRIDES[key], True
    return 0, False


def find_frame(step_name: str, runs: list[str],
               used: set[Path]) -> tuple[Path, str, int] | None:
    """Pick a frame that shows this step, and only this step.

    Almost every step name contains "cut", so rows are scored on the fraction
    of words they match and we take the middle of the best block, never t=0.
    """
    target = keywords(step_name)
    scored: list[tuple[float, str, dict]] = []
    for video_id in runs:
        rows = json.loads((OUTPUT_DIR / f"{video_id}_merged_steps.json").read_text())
        for row in rows:
            if row["start_sec"] <= 0:
                continue
            score = len(target & keywords(row["step"])) / len(target)
            if score:
                scored.append((score, video_id, row))
    if not scored:
        return None

    best = max(s for s, _, _ in scored)
    blocks: dict[str, list[dict]] = {}
    for score, video_id, row in scored:
        if score == best:
            blocks.setdefault(video_id, []).append(row)

    # Walk outward from the middle of the biggest block until we find a frame
    # no earlier step has already claimed.
    video_id, rows = max(blocks.items(), key=lambda kv: len(kv[1]))
    rows = sorted(rows, key=lambda r: r["start_sec"])
    mid = len(rows) // 2
    order = sorted(range(len(rows)), key=lambda i: abs(i - mid))
    for i in order:
        frame = FRAME_DIR / video_id / f"frame_{int(rows[i]['start_sec']):03d}s.jpg"
        if frame.exists() and frame not in used:
            used.add(frame)
            return frame, video_id, int(rows[i]["start_sec"])
    return None


MAX_CONDITION_FIGURES = 5


def condition_figures(sop: dict, runs: list[str],
                      used: set[Path]) -> dict[tuple[int, int], list[dict]]:
    """Find the frames where Molly hit a condition and changed what she did.

    Each such row is matched to the SOP decision point it belongs to by keyword
    overlap, keeping the best row per point and the points seen in most runs.
    """
    points = []
    for si, step in enumerate(sop["steps"], 1):
        for ti, task in enumerate(step["tasks"], 1):
            for di, dp in enumerate(task["decision_points"]):
                points.append((si, ti, di, dp))
    if not points:
        return {}

    # Best candidate row per decision point, scored on keyword overlap.
    best: dict[tuple[int, int, int], tuple[float, dict]] = {}
    for video_id in runs:
        path = OUTPUT_DIR / f"{video_id}_merged_steps.json"
        if not path.exists():
            continue
        for row in json.loads(path.read_text()):
            if not row.get("condition") or row.get("condition_response") != "adjusted":
                continue
            if row["start_sec"] <= 0:
                continue
            ts = int(row["start_sec"])
            frame = FRAME_DIR / video_id / f"frame_{ts:03d}s.jpg"
            if not frame.exists():
                continue
            for si, ti, di, dp in points:
                score = overlap(row["condition"], dp["condition"])
                if not score:
                    continue
                key = (si, ti, di)
                if key not in best or score > best[key][0]:
                    best[key] = (score, {"frame": frame, "video_id": video_id,
                                         "ts": ts, "condition": row["condition"],
                                         "observed_in_runs": dp["observed_in_runs"]})

    # Prefer the decision points seen in the most runs, then the cleanest match.
    ranked = sorted(best.items(),
                    key=lambda kv: (-kv[1][1]["observed_in_runs"], -kv[1][0]))
    figures: dict[tuple[int, int], list[dict]] = {}
    taken = 0
    for (si, ti, _di), (_score, fig) in ranked:
        if taken == MAX_CONDITION_FIGURES:
            break
        shift, overridden = override_for(fig["video_id"], fig["ts"])
        if overridden:
            if shift is None:          # reviewed and dropped
                continue
            shifted = FRAME_DIR / fig["video_id"] / f"frame_{fig['ts'] + shift:03d}s.jpg"
            if shifted.exists():
                fig["frame"], fig["ts"] = shifted, fig["ts"] + shift
        if fig["frame"] in used:  # no frame illustrates two things
            continue
        used.add(fig["frame"])
        figures.setdefault((si, ti), []).append(fig)
        taken += 1
    for group in figures.values():
        group.sort(key=lambda f: (f["video_id"], f["ts"]))
    return figures


def normalize(frame: Path, into: Path) -> Path:
    """Re-save a frame so python-docx will accept it.

    ffmpeg writes a valid JPEG that python-docx rejects, because it looks for a
    JFIF or Exif marker. Pillow writes a conventional header instead.
    """
    from PIL import Image
    out = into / frame.name
    Image.open(frame).convert("RGB").save(out, "JPEG", quality=90)
    return out


def page_number_footer(doc) -> None:
    """Center a page number in the footer.

    python-docx has no field API, so the field is assembled by hand.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    para = doc.sections[0].footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    for el, attrs, text in (("w:fldChar", {"w:fldCharType": "begin"}, None),
                            ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
                            ("w:fldChar", {"w:fldCharType": "end"}, None)):
        e = OxmlElement(el)
        for k, v in attrs.items():
            e.set(qn(k), v)
        if text:
            e.text = text
        run._r.append(e)


def build(sop: dict, group: str, with_images: bool) -> "Document":
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tmp = Path(tempfile.mkdtemp())
    used_frames: set[Path] = set()
    doc = Document()
    for name, size in (("Title", 24), ("Heading 1", 15), ("Heading 2", 12)):
        doc.styles[name].font.color.rgb = RGBColor(0x10, 0x11, 0x14)
        doc.styles[name].font.size = Pt(size)

    page_number_footer(doc)

    # Figures are numbered within their step - Figure 3.2 is the second figure
    # in step 3 - so a caption tells you where in the procedure it belongs
    # without having to scroll back and find it.
    figure_no: dict[int, int] = {}

    def add_figure(step_no: int, frame: Path, caption: str) -> None:
        figure_no[step_no] = figure_no.get(step_no, 0) + 1
        label = f"{step_no}.{figure_no[step_no]}"
        doc.add_picture(str(normalize(frame, tmp)), width=Inches(5.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = cap.add_run(f"Figure {label}. {caption}")
        c.font.size = Pt(8)
        c.font.color.rgb = RGBColor(0x8B, 0x8E, 0x98)

    # Chosen before the per-step frames, so the same picture cannot illustrate
    # both a condition and a step.
    cond_figs = (condition_figures(sop, GROUP_RUNS[group], used_frames)
                 if with_images else {})

    doc.add_heading(sop["title"], 0)

    intro = doc.add_paragraph()
    intro.add_run(
        "Derived procedure. ").bold = True
    intro.add_run(
        f"This document was not authored. It was synthesized from {len(GROUP_RUNS[group])} "
        "observed runs, and every step, action and decision point below was "
        "performed by an operator on video. The figure captions give the run and "
        "timestamp each illustration came from. Coverage notes such as "
        "“observed in 2 of 3 runs” indicate how consistently something was done, "
        "and are the reader’s guide to how firm each instruction is.")

    for i, step in enumerate(sop["steps"], 1):
        doc.add_heading(f"{i}. {step['step']}", 1)

        purpose = doc.add_paragraph()
        purpose.add_run(step["purpose"]).italic = True

        meta = f"Observed in {step['observed_in_runs']} of {len(GROUP_RUNS[group])} runs."
        if step["repeats"]:
            meta += f"  {step['repeats']}"
        note = doc.add_paragraph()
        run = note.add_run(meta)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x62, 0x65, 0x6E)

        # Step 1 is still setup, so its frames never showed the step itself.
        if with_images and i > 1:
            found = find_frame(step["step"], GROUP_RUNS[group], used_frames)
            if found:
                frame, video_id, ts = found
                add_figure(i, frame, f"{step['step']} — {video_id}, t={ts}s")

        for j, task in enumerate(step["tasks"], 1):
            doc.add_heading(f"{i}.{j} {task['task']}", 2)
            for k, action in enumerate(task["actions"], 1):
                para = doc.add_paragraph(f"{i}.{j}.{k} {action}")
                para.paragraph_format.left_indent = Inches(0.3)

            if task["decision_points"]:
                lead = doc.add_paragraph()
                lead.add_run("Decision points").bold = True

                table = doc.add_table(rows=1, cols=2)
                table.style = "Light Grid Accent 1"
                for cell, text in zip(table.rows[0].cells, ("If", "Then")):
                    cell.text = ""
                    cell.paragraphs[0].add_run(text).bold = True
                for dp in task["decision_points"]:
                    cells = table.add_row().cells
                    cells[0].text = dp["condition"]
                    cells[1].text = dp["response"]
                doc.add_paragraph()

            for fig in cond_figs.get((i, j), []):
                add_figure(i, fig["frame"],
                           f"{fig['condition']} — {fig['video_id']}, "
                           f"t={fig['ts']}s")

    return doc


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--group", default="molly", choices=list(GROUP_RUNS))
    parser.add_argument("--no-images", action="store_true")
    args = parser.parse_args()

    sop_path = SOP_DIR / f"sop_{args.group}.json"
    if not sop_path.exists():
        print(f"error: {sop_path.name} missing. Run "
              "pipeline/05_synthesize_sop.py first.", file=sys.stderr)
        return 1

    sop = json.loads(sop_path.read_text())
    doc = build(sop, args.group, not args.no_images)

    # Molly's SOP is what a customer is handed. The baseline SOP is comparison
    # material, so it stays with the analysis.
    out_dir = DELIVERABLE_DIR if args.group == "molly" else SOP_DIR
    out_dir.mkdir(exist_ok=True)
    # Named after the SOP's own title, so the file and the document cannot drift
    # apart. The colon is dropped because it is unsafe in a filename on macOS.
    out = out_dir / (sop["title"].replace("Standard Operating Procedure:", "SOP -") + ".docx")
    doc.save(out)

    steps = len(sop["steps"])
    decisions = sum(len(t["decision_points"]) for s in sop["steps"] for t in s["tasks"])
    figures = sum(1 for p in doc.inline_shapes)
    print(f"{steps} steps, {decisions} decision points, {figures} figures")
    print(f"-> {out.relative_to(REPO)}")

    # An override that stops matching is worse than none, because the bad frame
    # it was written to fix comes back without anyone noticing.
    if not args.no_images:
        # Only the overrides belonging to this group's runs, or every baseline
        # build would report Molly's overrides as stale.
        mine = {k for k in FIGURE_OVERRIDES if k[0] in GROUP_RUNS[args.group]}
        for video_id, ts in sorted(mine - _overrides_used):
            print(f"warning: FIGURE_OVERRIDES entry ({video_id}, {ts}) never matched "
                  f"a candidate frame - it may be stale", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
